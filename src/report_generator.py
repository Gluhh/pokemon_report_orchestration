"""Word report generation for Pokemon data."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

from src.config import Settings, get_settings
from src.database import PokemonRecord
from src.pokeapi_client import PokeAPIClient


def generate_pokemon_report(
    records: list[PokemonRecord],
    settings: Settings | None = None,
    client: PokeAPIClient | None = None,
) -> Path:
    settings = settings or get_settings()
    client = client or PokeAPIClient()

    settings.report_output_dir.mkdir(parents=True, exist_ok=True)
    settings.sprites_cache_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = settings.report_output_dir / f"pokemon_report_{timestamp}.docx"

    document = Document()
    document.add_heading("Complete Pokemon Report", level=0)
    document.add_paragraph(
        f"Report generated from PokeAPI data for all {len(records)} Pokemon, "
        "ordered by National Dex number."
    )
    document.add_paragraph(f"Generated at: {datetime.now():%Y-%m-%d %H:%M:%S}")
    document.add_page_break()

    for record in records:
        document.add_heading(f"#{record.id} - {record.name}", level=1)

        if record.sprite_url:
            sprite_path = settings.sprites_cache_dir / f"{record.id}_{record.name.lower().replace(' ', '_')}.png"
            try:
                client.download_sprite(record.sprite_url, str(sprite_path))
                document.add_picture(str(sprite_path), width=Inches(2.0))
            except Exception:
                document.add_paragraph(f"Sprite URL: {record.sprite_url}")

        document.add_paragraph(record.summary)
        document.add_paragraph("")

    document.save(report_path)
    return report_path
